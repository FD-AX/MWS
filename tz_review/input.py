from __future__ import annotations

from io import BytesIO
from pathlib import Path


class DocumentInputError(ValueError):
    """The uploaded document cannot be converted to reviewable text."""


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "cp1251"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentInputError("Не удалось определить кодировку текстового файла.")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:  # noqa: BLE001 - show a friendly upload error
        raise DocumentInputError("Не удалось прочитать PDF-файл.") from exc
    return "\n\n".join(page.strip() for page in pages if page.strip())


def _extract_docx(data: bytes) -> str:
    from docx import Document

    try:
        document = Document(BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - show a friendly upload error
        raise DocumentInputError("Не удалось прочитать DOCX-файл.") from exc

    blocks = [paragraph.text.strip() for paragraph in document.paragraphs
              if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def extract_text(filename: str, data: bytes) -> str:
    """Extract text from a supported document uploaded through the UI."""
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md"}:
        text = _decode_text(data)
    elif suffix == ".pdf":
        text = _extract_pdf(data)
    elif suffix == ".docx":
        text = _extract_docx(data)
    else:
        raise DocumentInputError(
            "Поддерживаются файлы PDF, DOCX, TXT и Markdown."
        )

    text = text.strip()
    if not text:
        raise DocumentInputError(
            "В документе не найден текст. Возможно, PDF состоит из сканов."
        )
    return text
